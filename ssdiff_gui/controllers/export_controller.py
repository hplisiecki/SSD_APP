"""Export functionality for SSD results.

Exports results to:
- CSV: scores, poles
- Word (.docx): regression table / pairwise comparison table, cluster tables, snippet tables
- PNG: PCA sweep plot
- JSON: config snapshot
"""

import shutil
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..models.project import Project, Run
from ..utils.export_settings import (
    get_export_setting,
    KEY_TOP_WORDS,
    KEY_CLUSTER_COS_BETA,
    KEY_CLUSTER_COHERENCE,
    KEY_CLUSTER_EXCERPT,
    KEY_CONT_ADJ_R2,
    KEY_CONT_F,
    KEY_CONT_BETA_NORM,
    KEY_CONT_DELTA,
    KEY_CONT_IQR,
    KEY_CONT_R,
    KEY_CROSS_N_COUNTS,
    KEY_CROSS_P_CORR,
    KEY_CROSS_COHENS_D,
    KEY_CROSS_CONTRAST_NORM,
)


# ------------------------------------------------------------------ #
#  Low-level docx helpers (adapted from reference code)
# ------------------------------------------------------------------ #

def _set_run_font(run, name="Times New Roman", size_pt=11, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_text(
    cell,
    text,
    *,
    align="left",
    bold=False,
    italic=False,
    size_pt=11,
    space_before_pt=0,
    space_after_pt=0,
    line_spacing=None,
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing

    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]

    r = p.add_run("" if text is None else str(text))
    _set_run_font(r, size_pt=size_pt, bold=bold, italic=italic)


def _set_cell_no_wrap(cell, no_wrap=True):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:noWrap"))
    if no_wrap:
        if existing is None:
            tcPr.append(OxmlElement("w:noWrap"))
    else:
        if existing is not None:
            tcPr.remove(existing)


def _clear_all_table_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge, edge_data in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for k, v in edge_data.items():
            element.set(qn(f"w:{k}"), str(v))


def _apply_apa_rules(table):
    """Apply APA-style horizontal rules: top/bottom of header, bottom of last row."""
    rule = {"val": "single", "sz": "8", "color": "000000"}
    for cell in table.rows[0].cells:
        _set_cell_border(cell, top=rule)
        _set_cell_border(cell, bottom=rule)
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom=rule)


def _setup_document(orientation="landscape", margins_in=1.0):
    """Create a new Document with standard margins and orientation."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(margins_in)
    section.bottom_margin = Inches(margins_in)
    section.left_margin = Inches(margins_in)
    section.right_margin = Inches(margins_in)

    if orientation.lower().startswith("land"):
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    return doc


def _add_table_caption(doc, table_number, title, appendix=False):
    """Add APA-style table caption (bold number + italic title)."""
    prefix = "A" if appendix else ""
    p1 = doc.add_paragraph()
    _set_run_font(p1.add_run(f"Table {prefix}{table_number}"), bold=True, size_pt=12)

    p2 = doc.add_paragraph()
    _set_run_font(p2.add_run(title), italic=True, size_pt=12)
    p2.paragraph_format.space_after = Pt(6)


def _add_table_note(doc, note):
    """Add APA-style table note."""
    if not note:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _set_run_font(p.add_run("Note. "), italic=True, size_pt=11)
    _set_run_font(p.add_run(note.replace("Note. ", "")), size_pt=11)


def _add_section_heading(doc, text, size_pt=14):
    """Add a bold heading paragraph."""
    p = doc.add_paragraph()
    _set_run_font(p.add_run(text), bold=True, size_pt=size_pt)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


# ------------------------------------------------------------------ #
#  Width helpers
# ------------------------------------------------------------------ #

def _scale_col_widths(base_widths: dict, selected_cols: list,
                      available_in: float) -> dict:
    """Return a new width dict scaled so selected columns fill available_in inches."""
    total = sum(base_widths.get(c, 1.0) for c in selected_cols)
    if total <= 0:
        return {c: base_widths.get(c, 1.0) for c in selected_cols}
    scale = available_in / total
    return {c: base_widths.get(c, 1.0) * scale for c in selected_cols}


# ------------------------------------------------------------------ #
#  ExportController
# ------------------------------------------------------------------ #

class ExportController:
    """Handles exporting SSD results to various formats."""

    def __init__(self, run: Run, project: Project):
        self.run = run
        self.project = project

    @property
    def _is_crossgroup(self) -> bool:
        return (
            self.run.results is not None
            and self.run.results.analysis_type == "crossgroup"
        )

    def export_all(self, output_dir: Path) -> Path:
        """Export all results to a directory."""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        results = self.run.results

        # 1. Scores CSV
        self.export_scores_csv(output_dir / "scores.csv")

        # 2. Poles (neighbors) CSV
        self.export_poles_csv(output_dir / "poles.csv")

        # 3. Main results table (Word)
        if self._is_crossgroup:
            self.export_pairwise_docx(output_dir / "pairwise_comparison.docx")
        else:
            self.export_regression_docx(output_dir / "regression.docx")

        # 4. Cluster tables (Word)
        if results.clusters_summary or (
            self._is_crossgroup and results.contrast_results
        ):
            self.export_clusters_docx(output_dir / "clusters.docx")

        # 5. Snippet tables (Word)
        has_cluster_snippets = results.cluster_snippets_pos or results.cluster_snippets_neg
        has_beta_snippets = results.beta_snippets_pos or results.beta_snippets_neg

        if self._is_crossgroup and results.contrast_results:
            # For crossgroup, check any contrast has snippets
            for cr in results.contrast_results.values():
                if cr.get("cluster_snippets_pos") or cr.get("cluster_snippets_neg"):
                    has_cluster_snippets = True
                if cr.get("beta_snippets_pos") or cr.get("beta_snippets_neg"):
                    has_beta_snippets = True

        if has_cluster_snippets:
            self.export_cluster_snippets_docx(output_dir / "snippets_cluster.docx")
        if has_beta_snippets:
            label = "snippets_contrast.docx" if self._is_crossgroup else "snippets_beta.docx"
            self.export_beta_snippets_docx(output_dir / label)

        # 6. Copy PCA sweep plot if exists (continuous only)
        if not self._is_crossgroup:
            pca_plot = self.run.run_path / "pca_sweep_sweep_plot.png"
            if pca_plot.exists():
                shutil.copy(pca_plot, output_dir / "pca_sweep.png")

        # 7. Config summary
        self.export_config_json(output_dir / "config.json")

        # 8. Human-readable hyperparameters
        self.export_hyperparameters_txt(output_dir / "hyperparameters.txt")

        return output_dir

    # ------------------------------------------------------------------ #
    #  CSV exports
    # ------------------------------------------------------------------ #

    def export_scores_csv(self, filepath: Path):
        """Export per-document scores to CSV."""
        results = self.run.results

        if self._is_crossgroup and results.contrast_results:
            # Merge all contrast scores with a "contrast" column
            all_dfs = []
            for key, cr in results.contrast_results.items():
                scores = cr.get("contrast_scores", [])
                if scores:
                    df = pd.DataFrame(scores)
                    df.insert(0, "contrast", key)
                    all_dfs.append(df)
            if all_dfs:
                merged = pd.concat(all_dfs, ignore_index=True)
                merged.to_csv(filepath, index=False)
        elif results.doc_scores:
            df = pd.DataFrame(results.doc_scores)
            df.to_csv(filepath, index=False)

    def export_poles_csv(self, filepath: Path):
        """Export positive and negative poles to CSV."""
        results = self.run.results

        if self._is_crossgroup and results.contrast_results:
            # All contrasts with a "contrast" column
            rows = []
            for key, cr in results.contrast_results.items():
                for entry in cr.get("pos_neighbors", []):
                    rows.append({
                        "contrast": key,
                        "side": "positive",
                        "rank": entry.get("rank", ""),
                        "word": entry.get("word", ""),
                        "cosine": entry.get("cos", ""),
                    })
                for entry in cr.get("neg_neighbors", []):
                    rows.append({
                        "contrast": key,
                        "side": "negative",
                        "rank": entry.get("rank", ""),
                        "word": entry.get("word", ""),
                        "cosine": entry.get("cos", ""),
                    })
            if rows:
                df = pd.DataFrame(rows)
                df.to_csv(filepath, index=False)
        else:
            rows = []
            for entry in results.pos_neighbors:
                rows.append({
                    "side": "positive",
                    "rank": entry.get("rank", ""),
                    "word": entry.get("word", ""),
                    "cosine": entry.get("cos", ""),
                })
            for entry in results.neg_neighbors:
                rows.append({
                    "side": "negative",
                    "rank": entry.get("rank", ""),
                    "word": entry.get("word", ""),
                    "cosine": entry.get("cos", ""),
                })

            df = pd.DataFrame(rows)
            df.to_csv(filepath, index=False)

    # ------------------------------------------------------------------ #
    #  Word doc: Regression table (continuous)
    # ------------------------------------------------------------------ #

    def export_regression_docx(self, filepath: Path):
        """Export regression results as an APA-style Word table."""
        results = self.run.results
        doc = _setup_document(orientation="landscape")

        # Determine DV name from dataset config
        dv_name = self.run.frozen_dataset_config.outcome_column or "Outcome"

        _add_table_caption(
            doc,
            table_number=1,
            title="Regression results predicting the quantitative metric from PCVs",
        )

        # Columns: DV (required), optional, p (required), optional…
        _cont_optional = {
            "Adj R\u00b2":      get_export_setting(KEY_CONT_ADJ_R2),
            "F":                get_export_setting(KEY_CONT_F),
            "\u03b2\u0302 norm": get_export_setting(KEY_CONT_BETA_NORM),
            "\u0394 per 0.1":   get_export_setting(KEY_CONT_DELTA),
            "IQR":              get_export_setting(KEY_CONT_IQR),
            "r":                get_export_setting(KEY_CONT_R),
        }
        _cont_required = {"DV", "p"}

        all_cols = ["DV", "Adj R\u00b2", "F", "p",
                    "\u03b2\u0302 norm", "\u0394 per 0.1", "IQR", "r"]
        all_values = [
            dv_name,
            f"{results.r2_adj:.2f}",
            f"{results.f_stat:.2f}",
            f"{results.f_pvalue:.2e}" if results.f_pvalue < 0.001 else f"{results.f_pvalue:.3f}",
            f"{results.beta_norm_stdCN:.2f}",
            f"{results.delta_per_0p10_raw:.2f}",
            f"{results.iqr_effect_raw:.2f}",
            f"{results.y_corr_pred:.2f}",
        ]
        cols   = [c for c in all_cols   if c in _cont_required or _cont_optional.get(c, True)]
        values = [v for c, v in zip(all_cols, all_values)
                  if c in _cont_required or _cont_optional.get(c, True)]

        _base_widths = {
            "DV": 2.0,
            "Adj R\u00b2": 0.85,
            "F": 0.85,
            "p": 0.85,
            "\u03b2\u0302 norm": 1.0,
            "\u0394 per 0.1": 1.0,
            "IQR": 0.85,
            "r": 0.85,
        }
        col_widths = _scale_col_widths(_base_widths, cols, 9.0)  # landscape = 9" usable

        table = doc.add_table(rows=1, cols=len(cols))
        table.autofit = False
        _clear_all_table_borders(table)

        # Header row
        hdr = table.rows[0].cells
        for j, col in enumerate(cols):
            hdr[j].width = Inches(col_widths[col])
            hdr[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_no_wrap(hdr[j], True)
            _set_cell_text(hdr[j], col, align="center", bold=True, size_pt=11)

        # Data row
        row_cells = table.add_row().cells
        for j, (col, val) in enumerate(zip(cols, values)):
            row_cells[j].width = Inches(col_widths[col])
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            align = "left" if col == "DV" else "center"
            _set_cell_text(row_cells[j], val, align=align, size_pt=11)

        _apply_apa_rules(table)

        _add_table_note(
            doc,
            "R\u00b2 and Adj. R\u00b2 indicate the variance explained by the "
            "semantic gradient model and its adjusted value accounting for sample "
            "size. F and p refer to the omnibus significance test of the model. "
            "\u03b2\u0302 norm denotes the standardized magnitude of the semantic "
            "gradient (norm of the direction vector in embedding space). "
            "\u0394 per 0.1 represents the expected change in semantic alignment "
            "associated with a 0.1 increase in the questionnaire score. "
            "IQR is the interquartile range of semantic scores within each corpus, "
            "reflecting the degree of semantic variability across participants. "
            "r is the correlation between SSD-based scores and questionnaire scores.",
        )

        doc.save(str(filepath))

    # ------------------------------------------------------------------ #
    #  Word doc: Pairwise comparison table (crossgroup)
    # ------------------------------------------------------------------ #

    def export_pairwise_docx(self, filepath: Path):
        """Export pairwise comparison results as an APA-style Word table."""
        results = self.run.results
        doc = _setup_document(orientation="landscape")

        _add_table_caption(
            doc,
            table_number=1,
            title="Pairwise comparison results from permutation-based group centroid tests",
        )

        _cross_optional = {
            "n_A":               get_export_setting(KEY_CROSS_N_COUNTS),
            "n_B":               get_export_setting(KEY_CROSS_N_COUNTS),
            "p (corr)":          get_export_setting(KEY_CROSS_P_CORR),
            "Cohen's d":         get_export_setting(KEY_CROSS_COHENS_D),
            "\u2016Contrast\u2016": get_export_setting(KEY_CROSS_CONTRAST_NORM),
        }
        _cross_required = {"Group A", "Group B", "Cos Distance", "p"}

        all_cols = ["Group A", "Group B", "n_A", "n_B",
                    "Cos Distance", "p", "p (corr)", "Cohen's d", "\u2016Contrast\u2016"]
        cols = [c for c in all_cols if c in _cross_required or _cross_optional.get(c, True)]

        _base_widths = {
            "Group A": 1.2,
            "Group B": 1.2,
            "n_A": 0.6,
            "n_B": 0.6,
            "Cos Distance": 1.0,
            "p": 0.85,
            "p (corr)": 0.85,
            "Cohen's d": 0.85,
            "\u2016Contrast\u2016": 1.0,
        }
        col_widths = _scale_col_widths(_base_widths, cols, 9.0)  # landscape = 9" usable

        table = doc.add_table(rows=1, cols=len(cols))
        table.autofit = False
        _clear_all_table_borders(table)

        # Header row
        hdr = table.rows[0].cells
        for j, col in enumerate(cols):
            hdr[j].width = Inches(col_widths[col])
            hdr[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_no_wrap(hdr[j], True)
            _set_cell_text(hdr[j], col, align="center", bold=True, size_pt=11)

        # Data rows
        for pw in results.pairwise_table or []:
            p_raw = pw.get("p_raw", 0)
            p_corr = pw.get("p_corrected", 0)

            all_values = [
                pw.get("group_A", ""),
                pw.get("group_B", ""),
                f"{pw.get('n_A', 0):,}",
                f"{pw.get('n_B', 0):,}",
                f"{pw.get('cosine_distance', 0):.4f}",
                f"{p_raw:.2e}" if p_raw < 0.001 else f"{p_raw:.3f}",
                f"{p_corr:.2e}" if p_corr < 0.001 else f"{p_corr:.3f}",
                f"{pw.get('cohens_d', 0):.3f}",
                f"{pw.get('contrast_norm', 0):.4f}",
            ]
            values = [v for c, v in zip(all_cols, all_values)
                      if c in _cross_required or _cross_optional.get(c, True)]

            row_cells = table.add_row().cells
            for j, (col, val) in enumerate(zip(cols, values)):
                row_cells[j].width = Inches(col_widths[col])
                row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                align = "left" if col in ("Group A", "Group B") else "center"
                _set_cell_text(row_cells[j], val, align=align, size_pt=11)

        _apply_apa_rules(table)

        # Omnibus note
        note_parts = []
        if results.omnibus_T is not None and results.omnibus_p is not None:
            omnibus_p_str = (
                f"{results.omnibus_p:.2e}" if results.omnibus_p < 0.001
                else f"{results.omnibus_p:.3f}"
            )
            note_parts.append(
                f"Omnibus permutation test: T = {results.omnibus_T:.4f}, "
                f"p = {omnibus_p_str} ({results.n_perm or 5000:,} permutations)."
            )
        note_parts.append(
            "Cos Distance is the cosine distance between group centroids. "
            "p (corr) reflects Bonferroni-corrected p-values. "
            "Cohen's d is calculated from the permutation distribution. "
            "\u2016Contrast\u2016 is the L2 norm of the contrast vector."
        )

        _add_table_note(doc, " ".join(note_parts))

        doc.save(str(filepath))

    # ------------------------------------------------------------------ #
    #  Word doc: Cluster tables
    # ------------------------------------------------------------------ #

    def export_clusters_docx(self, filepath: Path):
        """Export cluster summary tables as an APA-style Word document.

        For continuous: one table for positive clusters and one for negative.
        For crossgroup: one pair of tables per contrast.
        """
        results = self.run.results
        doc = _setup_document(orientation="portrait")

        if self._is_crossgroup and results.contrast_results:
            table_num = 1
            for i, (key, cr) in enumerate(results.contrast_results.items()):
                if i > 0:
                    doc.add_page_break()

                _add_section_heading(doc, f"Contrast: {key}")

                parts = key.split(" vs ")
                g1_name = parts[0] if len(parts) > 0 else "Group A"
                g2_name = parts[1] if len(parts) > 1 else "Group B"

                pos_clusters = [c for c in (cr.get("clusters_summary") or []) if c.get("side") == "pos"]
                neg_clusters = [c for c in (cr.get("clusters_summary") or []) if c.get("side") == "neg"]

                if pos_clusters:
                    self._write_cluster_table(
                        doc, pos_clusters, "pos", table_num,
                        f"Positive semantic clusters ({g1_name} direction)",
                        snippets_pool=cr.get("cluster_snippets_pos") or [],
                    )
                    table_num += 1

                if neg_clusters:
                    if pos_clusters:
                        doc.add_page_break()
                    self._write_cluster_table(
                        doc, neg_clusters, "neg", table_num,
                        f"Negative semantic clusters ({g2_name} direction)",
                        snippets_pool=cr.get("cluster_snippets_neg") or [],
                    )
                    table_num += 1
        else:
            if not results.clusters_summary:
                return

            pos_clusters = [c for c in results.clusters_summary if c.get("side") == "pos"]
            neg_clusters = [c for c in results.clusters_summary if c.get("side") == "neg"]

            table_num = 1
            if pos_clusters:
                self._write_cluster_table(
                    doc, pos_clusters, "pos", table_num,
                    "Positive semantic clusters (+\u03b2 \u2192 higher outcome)",
                )
                table_num += 1

            if neg_clusters:
                if pos_clusters:
                    doc.add_page_break()
                self._write_cluster_table(
                    doc, neg_clusters, "neg", table_num,
                    "Negative semantic clusters (\u2212\u03b2 \u2192 lower outcome)",
                )

        doc.save(str(filepath))

    def _write_cluster_table(self, doc, clusters, side, table_number, title,
                             snippets_pool=None):
        """Write a single cluster summary table into the document."""
        results = self.run.results

        _add_table_caption(doc, table_number=table_number, title=title)

        _cluster_optional = {
            "Cos \u03b2":              get_export_setting(KEY_CLUSTER_COS_BETA),
            "Coherence":               get_export_setting(KEY_CLUSTER_COHERENCE),
            "Representative\nExcerpt": get_export_setting(KEY_CLUSTER_EXCERPT),
        }
        _cluster_required = {"No.", "Size", "Top Words"}
        _top_words_limit  = get_export_setting(KEY_TOP_WORDS)

        all_cluster_cols = ["No.", "Size", "Cos \u03b2", "Coherence", "Top Words",
                            "Representative\nExcerpt"]
        cols = [c for c in all_cluster_cols
                if c in _cluster_required or _cluster_optional.get(c, True)]

        _base_cluster_widths = {
            "No.": 0.4,
            "Size": 0.5,
            "Cos \u03b2": 0.7,
            "Coherence": 1.0,
            "Top Words": 1.7,
            "Representative\nExcerpt": 2.2,
        }
        col_widths = _scale_col_widths(_base_cluster_widths, cols, 6.5)  # portrait = 6.5" usable
        no_wrap_except = {"Top Words", "Representative\nExcerpt"}
        excerpt_col = "Representative\nExcerpt"
        va = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Determine snippet pool
        if snippets_pool is None:
            snippets_pool = (
                results.cluster_snippets_pos if side == "pos"
                else results.cluster_snippets_neg
            ) or []

        rows = []
        for c in clusters:
            rank = c.get("cluster_rank", "")
            centroid_label = f"{side}_cluster_{rank}"

            # Find representative excerpt (first/best snippet for this cluster)
            cluster_snippets = [
                s for s in snippets_pool
                if s.get("centroid_label") == centroid_label
            ]
            excerpt = ""
            if cluster_snippets:
                excerpt = cluster_snippets[0].get("snippet_anchor", "")

            top_words = c.get("top_words", "")
            if _top_words_limit and top_words:
                parts = top_words.split(", ")
                if len(parts) > _top_words_limit:
                    top_words = ", ".join(parts[:_top_words_limit]) + " \u2026"

            rows.append({
                "No.": rank,
                "Size": c.get("size", ""),
                "Cos \u03b2": c.get("centroid_cos_beta", ""),
                "Coherence": c.get("coherence", ""),
                "Top Words": top_words,
                "Representative\nExcerpt": excerpt,
            })

        # Create table
        table = doc.add_table(rows=1, cols=len(cols))
        table.autofit = False
        _clear_all_table_borders(table)

        # Header
        hdr = table.rows[0].cells
        for j, col in enumerate(cols):
            hdr[j].width = Inches(col_widths[col])
            hdr[j].vertical_alignment = va
            if col not in no_wrap_except:
                _set_cell_no_wrap(hdr[j], True)
            _set_cell_text(hdr[j], col, align="center", bold=True, size_pt=11)

        # Body
        for row_data in rows:
            row_cells = table.add_row().cells
            for j, col in enumerate(cols):
                val = row_data[col]
                row_cells[j].width = Inches(col_widths[col])
                row_cells[j].vertical_alignment = va

                if col == excerpt_col:
                    _set_cell_text(
                        row_cells[j], val, align="left", size_pt=11,
                        space_after_pt=4, line_spacing=1.0,
                    )
                elif col in ("Cos \u03b2", "Coherence"):
                    formatted = f"{float(val):.2f}" if val != "" else ""
                    _set_cell_text(row_cells[j], formatted, align="center", size_pt=11)
                elif col in ("No.", "Size"):
                    _set_cell_text(row_cells[j], val, align="center", size_pt=11)
                else:
                    _set_cell_text(row_cells[j], val, align="left", size_pt=11)

        _apply_apa_rules(table)

    # ------------------------------------------------------------------ #
    #  Word doc: Snippet tables
    # ------------------------------------------------------------------ #

    def export_cluster_snippets_docx(self, filepath: Path):
        """Export cluster snippets as APA-style Word tables."""
        results = self.run.results
        doc = _setup_document(orientation="portrait")

        if self._is_crossgroup and results.contrast_results:
            self._export_crossgroup_cluster_snippets(doc, results)
        else:
            self._export_continuous_cluster_snippets(doc, results)

        doc.save(str(filepath))

    def _export_continuous_cluster_snippets(self, doc, results):
        """Export cluster snippets for continuous analysis."""
        dfs = []
        titles = []

        for side, pool, side_label in [
            ("pos", results.cluster_snippets_pos or [], "Positive"),
            ("neg", results.cluster_snippets_neg or [], "Negative"),
        ]:
            if not pool:
                continue

            clusters_seen = []
            groups = {}
            for s in pool:
                label = s.get("centroid_label", "unknown")
                if label not in groups:
                    groups[label] = []
                    clusters_seen.append(label)
                groups[label].append(s)

            for label in clusters_seen:
                snippets = groups[label]
                df = pd.DataFrame({
                    "Cosine Similarity": [s.get("cosine", 0) for s in snippets],
                    "Snippet Anchor": [s.get("snippet_anchor", "") for s in snippets],
                })
                dfs.append(df)
                cluster_num = label.rsplit("_", 1)[-1]
                titles.append(
                    f"{side_label} cluster snippet anchors "
                    f"({side_label.lower()} cluster no. {cluster_num})"
                )

        if dfs:
            self._write_snippet_tables(doc, dfs, titles)

    def _export_crossgroup_cluster_snippets(self, doc, results):
        """Export cluster snippets for crossgroup analysis — all contrasts."""
        first = True
        table_num = 1

        for key, cr in results.contrast_results.items():
            if not first:
                doc.add_page_break()
            first = False

            _add_section_heading(doc, f"Contrast: {key}")

            parts = key.split(" vs ")
            g1_name = parts[0] if len(parts) > 0 else "Group A"
            g2_name = parts[1] if len(parts) > 1 else "Group B"

            dfs = []
            titles = []

            for side_key, side_label in [
                ("cluster_snippets_pos", g1_name),
                ("cluster_snippets_neg", g2_name),
            ]:
                pool = cr.get(side_key) or []
                if not pool:
                    continue

                clusters_seen = []
                groups = {}
                for s in pool:
                    label = s.get("centroid_label", "unknown")
                    if label not in groups:
                        groups[label] = []
                        clusters_seen.append(label)
                    groups[label].append(s)

                for label in clusters_seen:
                    snippets = groups[label]
                    df = pd.DataFrame({
                        "Cosine Similarity": [s.get("cosine", 0) for s in snippets],
                        "Snippet Anchor": [s.get("snippet_anchor", "") for s in snippets],
                    })
                    dfs.append(df)
                    cluster_num = label.rsplit("_", 1)[-1]
                    titles.append(
                        f"{side_label} direction cluster snippet anchors "
                        f"(cluster no. {cluster_num})"
                    )

            if dfs:
                self._write_snippet_tables(doc, dfs, titles, start_num=table_num)
                table_num += len(dfs)

    def export_beta_snippets_docx(self, filepath: Path):
        """Export beta/contrast snippets as APA-style Word tables."""
        results = self.run.results
        doc = _setup_document(orientation="portrait")

        if self._is_crossgroup and results.contrast_results:
            self._export_crossgroup_beta_snippets(doc, results)
        else:
            self._export_continuous_beta_snippets(doc, results)

        doc.save(str(filepath))

    def _export_continuous_beta_snippets(self, doc, results):
        """Export beta snippets for continuous analysis."""
        dfs = []
        titles = []

        for pool, side_label in [
            (results.beta_snippets_pos or [], "Positive"),
            (results.beta_snippets_neg or [], "Negative"),
        ]:
            if not pool:
                continue

            df = pd.DataFrame({
                "Cosine Similarity": [s.get("cosine", 0) for s in pool],
                "Snippet Anchor": [s.get("snippet_anchor", "") for s in pool],
            })
            dfs.append(df)
            titles.append(
                f"{side_label} \u03b2-direction snippet anchors"
            )

        if dfs:
            self._write_snippet_tables(doc, dfs, titles)

    def _export_crossgroup_beta_snippets(self, doc, results):
        """Export contrast snippets for crossgroup analysis — all contrasts."""
        first = True
        table_num = 1

        for key, cr in results.contrast_results.items():
            if not first:
                doc.add_page_break()
            first = False

            _add_section_heading(doc, f"Contrast: {key}")

            parts = key.split(" vs ")
            g1_name = parts[0] if len(parts) > 0 else "Group A"
            g2_name = parts[1] if len(parts) > 1 else "Group B"

            dfs = []
            titles = []

            for side_key, side_label in [
                ("beta_snippets_pos", g1_name),
                ("beta_snippets_neg", g2_name),
            ]:
                pool = cr.get(side_key) or []
                if not pool:
                    continue

                df = pd.DataFrame({
                    "Cosine Similarity": [s.get("cosine", 0) for s in pool],
                    "Snippet Anchor": [s.get("snippet_anchor", "") for s in pool],
                })
                dfs.append(df)
                titles.append(
                    f"{side_label} direction contrast snippet anchors"
                )

            if dfs:
                self._write_snippet_tables(doc, dfs, titles, start_num=table_num)
                table_num += len(dfs)

    def _write_snippet_tables(self, doc, dfs, titles, start_num=1):
        """Write multiple snippet (cos, text) tables into a document."""
        cos_w_in = Inches(1.1)
        txt_w_in = Inches(5.7)
        va = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        col_cos = "Cosine Similarity"
        col_text = "Snippet Anchor"

        for i, (df, title) in enumerate(zip(dfs, titles)):
            if i > 0:
                doc.add_page_break()

            _add_table_caption(
                doc, table_number=start_num + i, title=title, appendix=True,
            )

            # Limit to first 10 snippets per table
            disp = df[[col_cos, col_text]].head(10).copy()
            disp[col_cos] = pd.to_numeric(disp[col_cos], errors="coerce").map(
                lambda x: "" if pd.isna(x) else f"{float(x):.2f}"
            )
            disp[col_text] = disp[col_text].fillna("").astype(str)

            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            _clear_all_table_borders(table)

            # Header
            hdr = table.rows[0].cells
            hdr[0].width = cos_w_in
            hdr[1].width = txt_w_in
            hdr[0].vertical_alignment = va
            hdr[1].vertical_alignment = va
            _set_cell_no_wrap(hdr[0], True)
            _set_cell_no_wrap(hdr[1], True)
            _set_cell_text(hdr[0], col_cos, align="center", bold=True, size_pt=11)
            _set_cell_text(hdr[1], col_text, align="center", bold=True, size_pt=11)
            table.rows[0].heading = True

            # Body
            for row_idx in range(len(disp)):
                row = table.add_row().cells
                row[0].width = cos_w_in
                row[1].width = txt_w_in
                row[0].vertical_alignment = va
                row[1].vertical_alignment = va

                _set_cell_text(row[0], disp.iloc[row_idx, 0], align="center", size_pt=11)
                _set_cell_text(
                    row[1], disp.iloc[row_idx, 1], align="left", size_pt=11,
                    space_after_pt=6, line_spacing=1,
                )

            _apply_apa_rules(table)

    # ------------------------------------------------------------------ #
    #  JSON config
    # ------------------------------------------------------------------ #

    def export_config_json(self, filepath: Path):
        """Export configuration as JSON."""
        import json

        config = {
            "run_id": self.run.run_id,
            "timestamp": self.run.timestamp.isoformat(),
            "concept_config": self.run.concept_config.to_dict(),
            "dataset_config": self.run.frozen_dataset_config.to_dict(),
            "spacy_config": self.run.frozen_spacy_config.to_dict(),
            "embedding_config": self.run.frozen_embedding_config.to_dict(),
            "hyperparameters": self.run.frozen_hyperparameters.to_dict(),
        }

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2, ensure_ascii=False)

    # ------------------------------------------------------------------ #
    #  Human-readable hyperparameters
    # ------------------------------------------------------------------ #

    def export_hyperparameters_txt(self, filepath: Path):
        """Export all hyperparameters as a human-readable text file."""
        run = self.run
        ds = run.frozen_dataset_config
        sp = run.frozen_spacy_config
        emb = run.frozen_embedding_config
        hp = run.frozen_hyperparameters
        cc = run.concept_config
        res = run.results

        lines = []
        lines.append("=" * 60)
        lines.append("SSD Analysis — Hyperparameters & Configuration")
        lines.append("=" * 60)
        lines.append(f"Run ID:      {run.run_id}")
        lines.append(f"Timestamp:   {run.timestamp.isoformat()}")
        lines.append("")

        # -- Dataset --
        lines.append("-" * 60)
        lines.append("DATASET")
        lines.append("-" * 60)
        lines.append(f"Data file:           {ds.csv_path}")
        lines.append(f"Text column:         {ds.text_column}")
        lines.append(f"ID column:           {ds.id_column or '(none)'}")
        lines.append(f"Analysis type:       {ds.analysis_type}")
        if ds.analysis_type == "continuous":
            lines.append(f"Outcome column:      {ds.outcome_column}")
        else:
            lines.append(f"Group column:        {ds.group_column}")
            lines.append(f"Permutations:        {ds.n_perm}")
        lines.append(f"Concept mode:        {ds.concept_mode}")
        lines.append(f"Total rows:          {ds.n_rows}")
        lines.append(f"Valid rows:          {ds.n_valid}")
        lines.append("")

        # -- Text Processing --
        lines.append("-" * 60)
        lines.append("TEXT PROCESSING (spaCy)")
        lines.append("-" * 60)
        lines.append(f"Language:            {sp.language}")
        lines.append(f"spaCy model:         {sp.model}")
        lines.append(f"Lemmatize:           {sp.lemmatize}")
        lines.append(f"Remove stopwords:    {sp.remove_stopwords}")
        lines.append(f"Docs processed:      {sp.n_docs_processed}")
        lines.append(f"Total tokens:        {sp.total_tokens}")
        lines.append(f"Mean words/doc (pre-stopword): {sp.mean_words_before_stopwords:.1f}")
        lines.append("")

        # -- Embeddings --
        lines.append("-" * 60)
        lines.append("WORD EMBEDDINGS")
        lines.append("-" * 60)
        lines.append(f"Model path:          {emb.model_path}")
        if emb.model_name:
            lines.append(f"Model name:          {emb.model_name}")
        lines.append(f"L2 normalize:        {emb.l2_normalize}")
        lines.append(f"ABTT enabled:        {emb.abtt_enabled}")
        lines.append(f"ABTT m:              {emb.abtt_m}")
        lines.append(f"Vocab size:          {emb.vocab_size}")
        lines.append(f"Embedding dim:       {emb.embedding_dim}")
        lines.append(f"Coverage:            {emb.coverage_pct:.1f}%")
        lines.append(f"OOV tokens:          {emb.n_oov}")
        lines.append("")

        # -- Concept Definition --
        lines.append("-" * 60)
        lines.append("CONCEPT DEFINITION")
        lines.append("-" * 60)
        lines.append(f"Mode:                {cc.mode}")
        if cc.mode == "lexicon" and cc.lexicon_tokens:
            tokens_str = ", ".join(sorted(cc.lexicon_tokens))
            lines.append(f"Lexicon tokens ({len(cc.lexicon_tokens)}): {tokens_str}")
        if cc.min_hits_per_doc is not None:
            lines.append(f"Min hits/doc:        {cc.min_hits_per_doc}")
        lines.append(f"Drop no-hit docs:    {cc.drop_no_hits}")
        if cc.mode == "fulldoc" and cc.stoplist:
            lines.append(f"Stoplist ({len(cc.stoplist)}): {', '.join(sorted(cc.stoplist))}")
        lines.append(f"Coverage:            {cc.coverage_pct:.1f}%")
        lines.append(f"Docs with hits:      {cc.n_docs_with_hits}")
        lines.append(f"Median hits/doc:     {cc.median_hits:.1f}")
        lines.append(f"Mean hits/doc:       {cc.mean_hits:.1f}")
        lines.append("")

        # -- Hyperparameters --
        lines.append("-" * 60)
        lines.append("HYPERPARAMETERS")
        lines.append("-" * 60)

        lines.append("")
        lines.append("PCA Dimensionality:")
        lines.append(f"  Mode:              {hp.n_pca_mode}")
        if hp.n_pca_mode == "manual":
            lines.append(f"  Manual k:          {hp.n_pca_manual}")
        lines.append(f"  Sweep k range:     {hp.sweep_k_min} - {hp.sweep_k_max} (step {hp.sweep_k_step})")
        lines.append(f"  Stability crit:    {hp.sweep_stability_criterion}")
        lines.append(f"  AUCK radius:       {hp.auck_radius}")
        lines.append(f"  Beta smooth window:{hp.beta_smooth_win}")
        lines.append(f"  Beta smooth kind:  {hp.beta_smooth_kind}")
        lines.append(f"  Weight by size:    {hp.weight_by_size}")
        if res and res.selected_k is not None:
            lines.append(f"  Selected k:        {res.selected_k}")
            lines.append(f"  Var explained:     {res.pca_var_explained:.4f}")

        lines.append("")
        lines.append("Context & Weighting:")
        lines.append(f"  Window size (+/-): {hp.context_window_size}")
        lines.append(f"  SIF parameter (a): {hp.sif_a}")

        lines.append("")
        lines.append("Model:")
        lines.append(f"  Use unit beta:     {hp.use_unit_beta}")
        lines.append(f"  L2 normalize docs: {hp.l2_normalize_docs}")

        lines.append("")
        lines.append("Clustering:")
        lines.append(f"  Top N neighbors:   {hp.clustering_topn}")
        lines.append(f"  Auto-select K:     {hp.clustering_k_auto}")
        lines.append(f"  K range:           {hp.clustering_k_min} - {hp.clustering_k_max}")
        lines.append(f"  Top words/cluster: {hp.clustering_top_words}")
        lines.append("")

        # -- Sample --
        if res:
            lines.append("-" * 60)
            lines.append("SAMPLE")
            lines.append("-" * 60)
            lines.append(f"N raw:               {res.n_raw}")
            lines.append(f"N kept:              {res.n_kept}")
            lines.append(f"N dropped:           {res.n_dropped}")
            lines.append("")

        lines.append("=" * 60)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")
